# -*- coding: utf-8 -*-
"""
Created on Thu Sep 29 10:32:49 2022

@author: ngoel16
"""

# importing packages 

from docx import Document


# read the word file 
resume_text = ""

with open("sample_resume.docx", "rb") as f:
    doc_resume = Document(f)
    para = doc_resume.paragraphs


# extract text 
for p in para:
    if p.text:
        resume_text += p.text    

tables = doc_resume.tables
for table in tables:
    doctbls=[]
    tbllist=[]
    rowlist=[]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            resume_text += cell.text + " "


# read the description 
with open("description.txt", "r") as fp:
    desc = fp.read()

print("Description: ")
print(desc)

# Text cleaning 
def text_cleaning(text):
    text = text.lower()  # to lower case 
    
    return text 

print("Resume: ")
print(text_cleaning(resume_text))

print("Description: ")
print(text_cleaning(desc))
# Take out tokens 


# remove stopwords

# count tokens 

# match the count of tokens with the description tokens 
